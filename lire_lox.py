import re
from docx import Document
import os
import platform

def nettoyer_texte_pour_XML(texte):
    return re.sub(r'[^\x20-\x7E]', '', texte)

def lire_contenu_docx(nom_fichier_docx):
    try:
        document = Document(nom_fichier_docx)
        for paragraphe in document.paragraphs:
            print(paragraphe.text)
    except Exception as e:
        print(f"Une erreur s'est produite lors de la lecture du fichier DOCX : {str(e)}")

def lire_et_ecrire_lox(nom_fichier_lox):
    try:
        encodages = ['UTF-8', 'UTF-16', 'UTF-16-le', 'UTF-16-be', 'UTF-32', 'UTF-32-le', 'UTF-32-be', 'windows-1250', 'windows-1251', 'windows-1252', 'ASCII', 
                     'Shift-JIS', 'euc-jp', 'big5', 'gb2312', 'ISO 8859-1', 'ISO 8859-2', 'ISO 8859-15','cp437', 'mac_roman' , 'koi8-r' ]

        lignes_lox = None
        for encodage in encodages:
            try:
                with open(nom_fichier_lox, 'r', encoding=encodage) as file:
                    lignes_lox = file.readlines()
                print(f"Décodage réussi avec l'encodage {encodage}")
                break
            except UnicodeDecodeError:
                print(f"Erreur de décodage avec l'encodage {encodage}, tentative avec un autre encodage.")

        if lignes_lox is None:
            raise Exception("Impossible de décoder le fichier avec les encodages disponibles.")

        nom_fichier_sortie = "resultat.docx"

        document = Document()
        for ligne in lignes_lox:
            ligne_nettoyee = nettoyer_texte_pour_XML(ligne.strip())
            if ligne_nettoyee:
                document.add_paragraph(ligne_nettoyee)  

        document.save(nom_fichier_sortie)
        print(f"Les données ont été écrites dans le fichier {nom_fichier_sortie}")

        systeme = platform.system()

        if systeme == "Windows":
                os.system(f'start winword "{nom_fichier_sortie}"')  
        elif systeme == "Darwin":
                os.system(f'open -a TextEdit "{nom_fichier_sortie}"')  
        elif systeme == "Linux":
                os.system(f'xdg-open "{nom_fichier_sortie}"')  
        else:
                print("Système d'exploitation non supporté pour l'ouverture automatique.")

        print(f"Vérification du contenu de {nom_fichier_sortie} avec python-docx...")
        lire_contenu_docx(nom_fichier_sortie)  

        print("Essayez d'ouvrir le fichier avec d'autres programmes si nécessaire.")
        
        print(f"Le fichier {nom_fichier_sortie} n'a pas été trouvé.")

    except FileNotFoundError:
        print(f"Le fichier {nom_fichier_lox} n'a pas été trouvé. Veuillez vérifier le chemin et le nom du fichier.")
    except Exception as e:
        print(f"Une erreur s'est produite : {str(e)}")


nom_fichier_lox = "nom de votre fichier"
lire_et_ecrire_lox(nom_fichier_lox)